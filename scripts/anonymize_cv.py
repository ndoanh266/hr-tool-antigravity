import os
import re
import sys
import argparse

# Try importing dependencies
try:
    import docx
except ImportError:
    docx = None

try:
    import win32com.client
except ImportError:
    win32com = None

def convert_japanese_year(text):
    # Regex to find Japanese eras: 平成 (Heisei), 昭和 (Showa), 令和 (Reiwa)
    heisei_pattern = re.compile(r'平成\s*(\d+)\s*年')
    showa_pattern = re.compile(r'昭和\s*(\d+)\s*年')
    reiwa_pattern = re.compile(r'令和\s*(\d+)\s*年')
    
    def repl_heisei(match):
        year = int(match.group(1))
        return f"{1988 + year} (平成{year}年)"
        
    def repl_showa(match):
        year = int(match.group(1))
        return f"{1925 + year} (昭和{year}年)"
        
    def repl_reiwa(match):
        year = int(match.group(1))
        return f"{2018 + year} (令和{year}年)"

    text = heisei_pattern.sub(repl_heisei, text)
    text = showa_pattern.sub(repl_showa, text)
    text = reiwa_pattern.sub(repl_reiwa, text)
    return text

def anonymize_text(text):
    # 1. Hide Phone numbers (standard Vietnamese mobile phone regex)
    phone_pattern = re.compile(r'\b(0[35789]\d{8}|02\d{8,9}|\+84\d{9,10})\b')
    # Replace intermediate digits to preserve visual format but hide details
    text = phone_pattern.sub("[ĐÃ ẨN SĐT]", text)
    
    # 2. Hide Emails
    email_pattern = re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b')
    text = email_pattern.sub("[ĐÃ ẨN EMAIL]", text)
    
    # 3. Handle Japanese years
    text = convert_japanese_year(text)
    
    return text

def anonymize_docx(file_path, output_path):
    if not docx:
        print("[LỖI] Thư viện python-docx chưa được cài đặt. Vui lòng chạy: pip install python-docx")
        return False
        
    doc = docx.Document(file_path)
    
    # Process all paragraphs
    for p in doc.paragraphs:
        if p.text:
            # We need to preserve run formatting if possible, but basic text replace is safer for runs
            for run in p.runs:
                if run.text:
                    run.text = anonymize_text(run.text)
                    
    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = anonymize_text(run.text)
                            
    doc.save(output_path)
    print(f"[OK] Đã che thông tin và lưu file Word mới tại: {output_path}")
    return True

def convert_to_pdf(docx_path, pdf_path):
    if not win32com:
        print("[LỖI] Thư viện pywin32 chưa được cài đặt. Không thể gọi MS Word để xuất PDF. Vui lòng chạy: pip install pywin32")
        return False
        
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 represents wdFormatPDF
        doc.Close()
        word.Quit()
        print(f"[OK] Đã xuất file PDF bảo mật tại: {pdf_path}")
        return True
    except Exception as e:
        print(f"[LỖI] Lỗi khi chuyển đổi PDF qua MS Word: {e}")
        return False

if __name__ == "__main__":
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser(description="Tự động bảo mật CV: Ẩn SĐT, Email, dịch năm Nhật và xuất PDF.")
    parser.add_argument("--input", required=True, help="Đường dẫn file CV gốc (DOCX)")
    parser.add_argument("--output_docx", required=True, help="Đường dẫn lưu file DOCX đã bảo mật")
    parser.add_argument("--output_pdf", required=True, help="Đường dẫn xuất file PDF")
    
    args = parser.parse_args()
    
    if anonymize_docx(args.input, args.output_docx):
        convert_to_pdf(args.output_docx, args.output_pdf)
